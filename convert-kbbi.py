
# coding: utf-8

# In[1]:


from docx import *
from docx.shared import RGBColor
import pandas as pd
import re


# In[12]:


def toStr(s):  
    str = " " 
    return (str.join(s))


# In[13]:


def clean(x):
    # hilangkan digit 
    temp1 = re.sub(r"\d","", x)

    # ilangkan tab 
    temp2 = re.sub(r"\t","", temp1)

    # hhilangkan spasi,tab, newline diakhir dan koma
    result = re.sub(r"\s$|,", "", temp2)
    
    return result


# In[14]:


def insertData(utama,turunan,similiar,kelas,panah,keterangan):
    x = [utama,turunan,similiar,kelas,panah,keterangan]
    # print(x)
    y = pd.DataFrame(data=[x],columns=['Kata Dasar','Kata Turunan','Kata Similar','Kelas Kata','Arah Panah','Keterangan'])
    
    return x,y


# In[15]:


def ambilKata(word):
    # jika ada tanda () merupakan similiar
    if '(' in word :
        x, y = re.sub(r"[/(].*?[/))=]","",word),re.sub(r"[()=]","",word)
        x, y = re.sub(r"\s$","",x),re.sub(r"\s$","",y)
        utama = x
        similar = y
        tempUtama = utama
                    
    else :
        word = re.sub(r"[()]|\s$|-\s","",word)
        utama = word
        similar = ""
        tempUtama = utama

    return utama,similar,tempUtama


# In[21]:


document = Document("test.docx")
# document = Document("file-kbbi-docx/KBBI-A.docx")

korpus = pd.DataFrame(columns=['Kata Dasar','Kata Turunan','Kata Similar','Kelas Kata','Arah Panah','Keterangan'])

kelasKata = ['n','v','a','adv','num','p','pron']
ket = ["Ab","Abr","Ach","Adm","Ag","akr","Anat","Antr","Ar","Arg","ark","Astron","Astr","Astrol","Bg","Bjr","Bio","Biol","Bl","Bld","Bot","Bt","Btk","Brk","cak","Cn","Dag","Dik","Dok","Dn","Dy","Ek","Ekg","El","Ent","Far","Fifol","Fil","Filol","Fils","Fis","Fir","Geo","Geog","Geol","Gra","Graf","Gy","Hd","hid","Hin","Hind","hor","Huk","Hut","Isl","Ikn","Ing","Isi","Isl","Jak","Jb","Jk","Jp","Jw","Jyw","k","Kal","Kap","kas","Kat","kep","Kes","Keu","Kim","kl","kp","Kris","Kom","Komp","lesl","Lay","Ling","Lis","lt","Man","Mat","Md","Mdr","Mek","Mes","Met","Mil","Mk","Mks","Mn","Mu","Mus","ok","Olr","on","Opt","Org","Orl","Pem","Pet","Plb","Pol","Psi","Pr","Prot","Publ","Sas","sb","sd","sel","sj","Sen","Sos","Skr","Skt","Sng","Sos","Stat","Tan","tbl","tas","Tek","terb","Terb","Tern","tld","Tns","Us","zat","Zool"]
kosong = ["", " " , "  "]
tempTurunan =""
tempUtama = ""

for paragraph in document.paragraphs:
    
    utama = ""
    turunan = ""
    similar = ""
    kelas = ""
    panah = ""
    keterangan = ""

    tempKelas = []
    tempKeterangan = []
#     tempUtama = ""
    
    
    index = 0
    for run in paragraph.runs:

        # teks warna merah (sublema)
        if ((run.font.color.rgb == RGBColor(255,0,0))):
#             print(run.text+","+str(index))
            
            # sublema
            if(run.bold):
                
                run.text = clean(run.text)
                
                if run.text not in kosong :
            
                    if utama != "" or turunan != "" or kelas != "" :
                        
                        # insert
                        x = insertData(utama,turunan,similar,kelas,panah,keterangan)
                        print(x[0])
                        korpus  = korpus.append([x[1]],ignore_index=True)
                        utama = ""
                    
                    # menangani jika -- ga bold
                    if index == 1 and re.search("--", paragraph.runs[index-1].text):

                        run.text = tempUtama+" "+run.text
                            
                        result = ambilKata(run.text)
                            
                        turunan     = result[0]
                        similar  = result[1]
                        kelas = ""
                        keterangan = ""
                        tempKeterangan = []
                        kelas = ""
                        
                    # menangani jika ~ ga bold
                    elif index == 1 and re.search("~", paragraph.runs[index-1].text):
#                         print(paragraph.runs[index+1].text)
#                         print(run.text)
                        run.text = tempTurunan+" "+run.text
                            
                        result = ambilKata(run.text)
                            
                        turunan     = result[0]
                        similar  = result[1]
                        kelas = ""
                        keterangan = ""
                        
                    else:
                        if re.search("--", run.text):
                            # ganti -- dengan lema (sebelumnya/tempUtama)
                            run.text = re.sub(r"--",tempUtama,run.text)

                            result = ambilKata(run.text)

                            turunan     = result[0]
                            similar  = result[1]
                            kelas = ""
                            keterangan = ""
                            tempKeterangan = []
                        
                        elif re.search("~", run.text):
#                             print(run.text)
                            # ganti ~ dengan sublema (sebelumnya/tempTurunan)
                            run.text = re.sub(r"~",tempTurunan,run.text)
                                
                            result = ambilKata(run.text)

                            turunan     = result[0]
                            similar  = result[1]
                            kelas = ""
                            keterangan = ""
                            tempKeterangan = []
                            
                        else:
                            result = ambilKata(run.text)
                        
                            turunan     = result[0]
                            similar    = result[1]
                            tempTurunan = result[2]
                            kelas = ""
                            keterangan = ""
                            tempKeterangan = []

                        
            elif(run.italic):

                run.text = clean(run.text)
                
                x = run.text.split(" ")

                # keterangan
                for y in x:
                    
                    if (y in kelasKata or y in ket) and y not in tempKeterangan:
                        tempKeterangan.append(y)
                        
                if tempKeterangan != "":
                    for temp in tempKeterangan :
                        if temp in ket:
                            keterangan = " ".join(tempKeterangan)
                
                # kelas kata
                for y in x:
                    
                    if y in kelasKata and y != kelas:
                        if kelas != "":
                            # insert
                            x = insertData(utama,turunan,similar,kelas,panah,keterangan)
                            print(x[0])
                            korpus  = korpus.append([x[1]],ignore_index=True)
                        kelas = y  
        
        # teks warna hijau (similar)
        elif ((run.font.color.rgb == RGBColor(0,176,80))):
#             print(run.text)
            if(run.bold):
                
                run.text = clean(run.text)
                
                if run.text not in kosong and re.search("\w", run.text) :
                    
                    similar = re.sub(r"[()=]","",run.text)
        
        # teks bukan warna merah (lema)
        else:
            
            # lema
            if(run.bold):
#                 print(run.text)
                run.text = clean(run.text)
                
                if run.text not in kosong and re.search("\w", run.text) :
                    
#                     # tanda panah
#                     temp = []
#                     if (re.search("‹", run.text)) :
# #                         print(run.text)
#                         tempPanah = paragraph.runs[index+2].text
#                         panah = tempPanah
                                     
#                         tempUtama = paragraph.runs[index-1].text
#                         utama = tempUtama
                        
#                     else :
                    
#                         if utama != "" or turunan!= "" or kelas != "":
#                             # insert
#                             x = insertData(utama,turunan,similar,kelas,panah,keterangan)
#                             print(x[0])
#                             korpus  = korpus.append([x[1]],ignore_index=True)

#                             turunan = ""

#                         result = ambilKata(run.text)

#                         utama     = result[0]
#                         similar  = result[1]
#                         tempUtama = result[2]
#                         kelas = ""
#                         keterangan = ""
#                         tempKeterangan = []
                    
                    if utama != "" or turunan!= "" or kelas != "":
                        # insert
                        x = insertData(utama,turunan,similar,kelas,panah,keterangan)
                        print(x[0])
                        korpus  = korpus.append([x[1]],ignore_index=True)

                        turunan = ""

                    result = ambilKata(run.text)

                    utama     = result[0]
                    similar  = result[1]
                    tempUtama = result[2]
                    kelas = ""
                    keterangan = ""
                    tempKeterangan = []
                
            elif(run.italic):

                run.text = clean(run.text)
                
                x = run.text.split(" ")

                # keterangan
                for y in x:
                    
                    if (y in kelasKata or y in ket) and y not in tempKeterangan:
                        tempKeterangan.append(y)
                        
                if tempKeterangan != "":
                    for temp in tempKeterangan :
                        if temp in ket:
                            keterangan = " ".join(tempKeterangan)
                
                # kelas kata
                for y in x:
                    
                    if y in kelasKata and y != kelas:
                        if kelas != "":
                            # insert
                            x = insertData(utama,turunan,similar,kelas,panah,keterangan)
                            print(x[0])
                            korpus  = korpus.append([x[1]],ignore_index=True)
                        kelas = y
            
            # tanda panah
            else:
                run.text = clean(run.text)
#                 print(run.text)
                
                if run.text not in kosong :
                    
                    # tanda panah
                    temp = []
                    if (re.search("‹", run.text)) :
                        
                        if re.search("\d", paragraph.runs[index+2].text):
                            tempPanah = paragraph.runs[index+3].text
                            panah = tempPanah
                        else :
                            tempPanah = paragraph.runs[index+2].text
                            panah = tempPanah
                                     
                        tempUtama = paragraph.runs[index-1].text
                        utama = tempUtama
                    
        index += 1
    
    # insert
    if utama not in kosong or turunan not in kosong or similar not in kosong or kelas not in kosong or panah not in kosong or keterangan not in kosong :
        x = insertData(utama,turunan,similar,kelas,panah,keterangan)
        print(x[0])
        korpus  = korpus.append([x[1]],ignore_index=True)
    


# In[11]:


korpus.to_excel (r'file-kbbi-excel/hasil-convert.xlsx', index = False, header=True)

